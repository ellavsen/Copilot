"""Build the FAISS index from a document corpus.

The original version could only read from three hard-coded Google Drive folder
URLs using a service-account key committed next to the code. That made the
project impossible for anyone else to run and leaked credentials. Here:

* the default source is a **local folder** (``data/corpus/``), so anyone who
  clones the repo can drop in PDFs and build an index;
* Google Drive is an optional source, configured entirely through environment
  variables, with no folder ids or keys in the source tree;
* document parsing and the embedding stack are optional extras, so importing
  this module never drags in torch.
"""

from __future__ import annotations

import logging
import re
import shutil
from collections.abc import Iterable, Iterator
from dataclasses import dataclass, field
from pathlib import Path

from indlab.config import Settings, get_settings

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {".md", ".txt", ".rst", ".markdown"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".pdf", ".docx"}


@dataclass(slots=True)
class CorpusDocument:
    """A source document, before chunking."""

    name: str
    text: str
    path: list[str] = field(default_factory=list)

    @property
    def path_str(self) -> str:
        return "/".join(self.path)


class MissingDependencyError(RuntimeError):
    pass


# ─────────────────────────────────────────────────────────────────────
# Local folder source (default)
# ─────────────────────────────────────────────────────────────────────
def _read_pdf(file: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise MissingDependencyError(
            'PDF support needs the ingest extra: pip install -e ".[ingest]"'
        ) from exc
    with fitz.open(file) as document:
        return "".join(page.get_text() for page in document)


def _read_docx(file: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise MissingDependencyError(
            'DOCX support needs the ingest extra: pip install -e ".[ingest]"'
        ) from exc
    return "\n".join(paragraph.text for paragraph in DocxDocument(str(file)).paragraphs)


def read_document(file: Path) -> str:
    """Extract plain text from one supported file."""
    suffix = file.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return file.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf(file)
    if suffix == ".docx":
        return _read_docx(file)
    raise ValueError(f"Unsupported file type: {file.suffix}")


def load_local_corpus(root: Path) -> Iterator[CorpusDocument]:
    """Walk ``root`` recursively and yield every readable document."""
    if not root.exists():
        log.warning("Corpus folder %s does not exist", root)
        return
    for file in sorted(root.rglob("*")):
        if not file.is_file() or file.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        try:
            text = read_document(file)
        except (MissingDependencyError, ValueError, OSError) as exc:
            log.warning("Skipping %s: %s", file.name, exc)
            continue
        if not text.strip():
            continue
        yield CorpusDocument(
            name=file.name,
            text=text,
            path=list(file.relative_to(root).parent.parts),
        )


# ─────────────────────────────────────────────────────────────────────
# Google Drive source (optional)
# ─────────────────────────────────────────────────────────────────────
def load_gdrive_corpus(
    folder_ids: Iterable[str], credentials_file: str
) -> Iterator[CorpusDocument]:
    """Yield documents from Google Drive folders.

    Requires the ``gdrive`` extra and a service-account key supplied at
    runtime. Never commit that key: point ``GDRIVE_SERVICE_ACCOUNT_FILE`` at a
    path outside the repository, or load it from a secret manager.
    """
    try:
        import io

        from google.oauth2 import service_account
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise MissingDependencyError(
            'Google Drive support needs: pip install -e ".[gdrive]"'
        ) from exc

    scopes = [
        "https://www.googleapis.com/auth/drive.readonly",
        "https://www.googleapis.com/auth/documents.readonly",
    ]
    credentials = service_account.Credentials.from_service_account_file(
        credentials_file, scopes=scopes
    )
    drive = build("drive", "v3", credentials=credentials, cache_discovery=False)
    docs = build("docs", "v1", credentials=credentials, cache_discovery=False)

    def download(file_id: str) -> bytes:
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, drive.files().get_media(fileId=file_id))
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buffer.getvalue()

    def walk(folder_id: str, parents: list[str]) -> Iterator[dict]:
        page_token = None
        while True:
            response = (
                drive.files()
                .list(
                    q=f"'{folder_id}' in parents and trashed=false",
                    fields="nextPageToken, files(id, name, mimeType)",
                    pageSize=1000,
                    pageToken=page_token,
                )
                .execute()
            )
            for item in response.get("files", []):
                if item["mimeType"] == "application/vnd.google-apps.folder":
                    yield from walk(item["id"], [*parents, item["name"]])
                else:
                    yield {**item, "path": parents}
            page_token = response.get("nextPageToken")
            if not page_token:
                break

    for folder_id in folder_ids:
        root_name = drive.files().get(fileId=folder_id, fields="name").execute()["name"]
        for item in walk(folder_id, [root_name]):
            name, file_id, mime = item["name"], item["id"], item["mimeType"]
            try:
                if mime == "application/vnd.google-apps.document":
                    document = docs.documents().get(documentId=file_id).execute()
                    text = "".join(
                        element.get("textRun", {}).get("content", "")
                        for block in document.get("body", {}).get("content", [])
                        for element in block.get("paragraph", {}).get("elements", [])
                    )
                elif name.lower().endswith(".pdf") or mime.endswith("pdf"):
                    import fitz

                    with fitz.open(stream=download(file_id), filetype="pdf") as pdf:
                        text = "".join(page.get_text() for page in pdf)
                elif name.lower().endswith(".docx") or "wordprocessingml" in mime:
                    import io as _io

                    from docx import Document as DocxDocument

                    text = "\n".join(
                        paragraph.text
                        for paragraph in DocxDocument(_io.BytesIO(download(file_id))).paragraphs
                    )
                elif name.lower().endswith((".md", ".txt")):
                    text = download(file_id).decode("utf-8", errors="ignore")
                else:
                    continue
            except Exception as exc:
                log.warning("Skipping %s: %s", name, exc)
                continue
            if text.strip():
                yield CorpusDocument(name=name, text=text, path=item["path"])


# ─────────────────────────────────────────────────────────────────────
# Chunking and index building
# ─────────────────────────────────────────────────────────────────────
def chunk_documents(documents: Iterable[CorpusDocument], settings: Settings) -> list:
    """Split documents into overlapping chunks carrying source metadata."""
    try:
        from langchain_core.documents import Document
        from langchain_text_splitters import RecursiveCharacterTextSplitter
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise MissingDependencyError('Chunking needs: pip install -e ".[rag]"') from exc

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = []
    for document in documents:
        normalised = re.sub(r"\s+", " ", document.text).strip()
        if not normalised:
            continue
        for index, piece in enumerate(splitter.split_text(normalised)):
            chunks.append(
                Document(
                    page_content=piece,
                    metadata={
                        "file_name": document.name,
                        "folders": document.path,
                        "path_str": document.path_str or document.name,
                        "chunk": index,
                    },
                )
            )
    return chunks


def build_index(documents: Iterable[CorpusDocument], settings: Settings | None = None) -> int:
    """Embed the corpus and write a fresh FAISS index. Returns chunk count."""
    settings = settings or get_settings()
    try:
        from langchain_community.vectorstores import FAISS
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise MissingDependencyError('Index building needs: pip install -e ".[rag]"') from exc

    chunks = chunk_documents(documents, settings)
    if not chunks:
        raise RuntimeError(
            f"No readable documents found. Put PDF/DOCX/MD files into {settings.corpus_dir} "
            "or configure GDRIVE_FOLDER_IDS."
        )

    log.info("Embedding %s chunks with %s …", len(chunks), settings.embedding_model)
    embeddings = HuggingFaceEmbeddings(
        model_name=settings.embedding_model,
        model_kwargs={"device": settings.embedding_device},
        encode_kwargs={"normalize_embeddings": True, "batch_size": 32},
    )
    store = FAISS.from_documents(chunks, embeddings)

    if settings.vector_db_path.exists():
        shutil.rmtree(settings.vector_db_path)
    settings.vector_db_path.parent.mkdir(parents=True, exist_ok=True)
    store.save_local(str(settings.vector_db_path))
    log.info("Index written to %s (%s chunks)", settings.vector_db_path, len(chunks))
    return len(chunks)


def ingest(
    *,
    corpus_dir: Path | None = None,
    gdrive_folder_ids: Iterable[str] | None = None,
    gdrive_credentials: str | None = None,
    settings: Settings | None = None,
) -> int:
    """Collect documents from the configured sources and rebuild the index."""
    settings = settings or get_settings()
    documents: list[CorpusDocument] = list(load_local_corpus(corpus_dir or settings.corpus_dir))
    log.info("Local corpus: %s documents", len(documents))

    if gdrive_folder_ids and gdrive_credentials:
        drive_documents = list(load_gdrive_corpus(gdrive_folder_ids, gdrive_credentials))
        log.info("Google Drive: %s documents", len(drive_documents))
        documents.extend(drive_documents)

    return build_index(documents, settings)
